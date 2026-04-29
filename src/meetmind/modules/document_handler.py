"""
Document handling using python-docx
"""
from typing import Optional, List
import logging
from pathlib import Path

logger = logging.getLogger(__name__)

class DocumentHandler:
    """Handle document creation and manipulation with python-docx."""

    def __init__(self, template_dir: Optional[str] = None):
        self.template_dir = Path(template_dir) if template_dir else Path(".")
        try:
            from docx import Document
            self.Document = Document
        except ImportError:
            logger.error("python-docx library not installed")
            self.Document = None

    def create_meeting_minutes(
        self,
        title: str,
        attendees: List[str],
        agenda: List[str],
        notes: str,
        action_items: List[str],
    ) -> Optional[str]:
        if not self.Document:
            logger.error("python-docx not available")
            return None

        try:
            output_dir = self.template_dir
            output_dir.mkdir(parents=True, exist_ok=True)
            doc = self.Document()

            doc.add_heading(f"Meeting Minutes: {title}", 0)
            doc.add_heading("Attendees", level=1)
            for attendee in attendees:
                doc.add_paragraph(attendee, style="List Bullet")

            doc.add_heading("Agenda", level=1)
            for item in agenda:
                doc.add_paragraph(item, style="List Number")

            doc.add_heading("Meeting Notes", level=1)
            doc.add_paragraph(notes or "No notes provided.")

            doc.add_heading("Action Items", level=1)
            for item in action_items:
                doc.add_paragraph(item, style="List Bullet")

            output_path = output_dir / f"{title.strip().replace(' ', '_')}_minutes.docx"
            doc.save(output_path)
            logger.info(f"Meeting minutes created: {output_path}")
            return str(output_path)
        except Exception as exc:
            logger.error(f"Error creating meeting minutes: {exc}")
            return None
